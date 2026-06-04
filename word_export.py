import re
from datetime import datetime
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt


def normalize_markdown_tables(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    output = []
    index = 0

    while index < len(lines):
        line = lines[index].replace("｜", "|")
        stripped = line.strip()

        if not (stripped.startswith("|") and stripped.endswith("|")):
            output.append(lines[index])
            index += 1
            continue

        table_lines = []
        while index < len(lines):
            candidate = lines[index].replace("｜", "|").strip()
            if candidate.startswith("|") and candidate.endswith("|"):
                table_lines.append(candidate)
                index += 1
            else:
                break

        if len(table_lines) > 1 and "---" not in table_lines[1]:
            columns = len(table_lines[0].strip("|").split("|"))
            table_lines.insert(1, "|" + "|".join(["---"] * columns) + "|")

        output.extend(table_lines)

    return "\n".join(output)


def markdown_to_html(markdown_text: str) -> str:
    markdown = MarkdownIt("gfm-like", {"html": False, "linkify": True})
    return markdown.render(normalize_markdown_tables(markdown_text))


def _set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    properties.append(shade)


def _set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D9E2EC")
        borders.append(border)
    table._tbl.tblPr.append(borders)


def _add_table(document: Document, rows: list[list[dict]]):
    if not rows:
        return

    columns = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)

    for column in range(columns):
        data = rows[0][column] if column < len(rows[0]) else {"text": ""}
        _set_cell_text(table.rows[0].cells[column], data.get("text", ""), bold=True)
        _shade_cell(table.rows[0].cells[column], "EAF2F8")

    for row in rows[1:]:
        cells = table.add_row().cells
        for column in range(columns):
            data = row[column] if column < len(row) else {"text": ""}
            _set_cell_text(cells[column], data.get("text", ""))

    document.add_paragraph()


def _style_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12.5, "1F4E79"),
        ("Heading 3", 11.5, "335C67"),
    ]:
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


class HtmlToDocxParser(HTMLParser):
    def __init__(self, document: Document):
        super().__init__()
        self.document = document
        self.paragraph = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.lists = []
        self.in_table = False
        self.table_rows = []
        self.row = None
        self.cell = None

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in {"strong", "b"}:
            self.bold = True
        elif tag in {"em", "i"}:
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag in {"h1", "h2", "h3"}:
            self.paragraph = self.document.add_heading("", level=int(tag[1]))
        elif tag in {"p", "div"} and not self.in_table:
            self.paragraph = self.document.add_paragraph()
        elif tag == "br" and self.paragraph is not None:
            self.paragraph.add_run().add_break()
        elif tag in {"ul", "ol"}:
            self.lists.append(tag)
        elif tag == "li":
            style = "List Number" if self.lists and self.lists[-1] == "ol" else "List Bullet"
            self.paragraph = self.document.add_paragraph(style=style)
        elif tag == "table":
            self.in_table = True
            self.table_rows = []
        elif tag == "tr":
            self.row = []
        elif tag in {"th", "td"}:
            self.cell = {"text": "", "header": tag == "th"}

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {"strong", "b"}:
            self.bold = False
        elif tag in {"em", "i"}:
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag in {"h1", "h2", "h3", "p", "div", "li"}:
            self.paragraph = None
        elif tag in {"ul", "ol"} and self.lists:
            self.lists.pop()
        elif tag in {"th", "td"} and self.row is not None and self.cell is not None:
            self.cell["text"] = re.sub(r"\s+", " ", self.cell["text"]).strip()
            self.row.append(self.cell)
            self.cell = None
        elif tag == "tr" and self.row is not None:
            self.table_rows.append(self.row)
            self.row = None
        elif tag == "table":
            _add_table(self.document, self.table_rows)
            self.in_table = False

    def handle_data(self, data):
        if self.in_table and self.cell is not None:
            self.cell["text"] += data
            return
        if not data.strip():
            return
        if self.paragraph is None:
            self.paragraph = self.document.add_paragraph()
        run = self.paragraph.add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def html_to_docx_bytes(html_text: str) -> bytes:
    document = Document()
    _style_document(document)
    document.core_properties.title = "项目简报"
    document.core_properties.author = "广东九丰能源项目简报AI Agent"
    document.core_properties.created = datetime.now()

    parser = HtmlToDocxParser(document)
    parser.feed(html_text or "<h1>项目简报</h1><p>暂无内容。</p>")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def word_filename() -> str:
    return f"项目简报_{datetime.now().strftime('%Y%m%d')}.docx"
